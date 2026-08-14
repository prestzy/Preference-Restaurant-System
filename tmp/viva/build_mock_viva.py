from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT_DIR = ROOT / "output" / "viva"
OUTPUT_PATH = OUTPUT_DIR / "Personalized-Restaurant-Ordering-System-Mock-Viva-QA.docx"


# compact_reference_guide preset tokens.
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
MUTED = "5D6775"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
RISK_RED = "9B1C1C"
LIGHT_RED = "FDECEC"
ORANGE = "F97316"  # Named project-accent override, used only for labels.
WHITE = "FFFFFF"


SECTIONS = [
    (
        "1. Research Framing and Contribution",
        [
            {
                "q": "Be honest: is this not just a QR menu with a few if-statements?",
                "a": "The QR menu itself is not the research contribution. The contribution is an artefact-oriented integration of a complete ordering workflow with a transparent recommender designed for a small, sparse-data restaurant. The system separates eligibility, ranking, and evidence confidence; combines explicit preferences with co-order evidence; and provides controlled, non-destructive evaluation tools. I use established recommendation ideas and do not claim that the web menu or the scoring rules are a new branch of artificial intelligence.",
                "e": "Report Sections 5.2-5.3 describe the contribution as application-focused and artefact-oriented, not a new mathematical model.",
                "t": "Do not answer, 'It is novel because it uses AI.' The system is deterministic and rule-based.",
            },
            {
                "q": "What exactly is novel if content-based and collaborative filtering already exist?",
                "a": "The novelty is contextual and design-oriented rather than algorithmic. I adapted established techniques to a limited-data, single-restaurant setting and connected them to hard restrictions, adaptive evidence weighting, plain-language explanations, order completion, and repeatable evaluation. The strongest claim is that this design demonstrates a practical and inspectable way to embed lightweight personalization in an ordering system. It is not a claim that I invented content-based or collaborative filtering.",
                "e": "The report explicitly states that established techniques were adapted into a working solution rather than used to propose a new mathematical model.",
                "t": "Avoid claiming algorithmic novelty. Examiners can disprove that immediately from the literature.",
            },
            {
                "q": "How do your research questions actually connect to what you built?",
                "a": "Each research question maps to an observable mechanism and experiment. RQ1 tests whether liked ingredients move dishes upward and disliked ingredients enforce exclusion. RQ2 tests whether increasing co-order evidence changes support, confidence, lift, and rank. RQ3 compares ingredient-only, co-order-only, and fixed-hybrid recovery of a hidden dish. The operational system implements those same signals, while the administrator tester isolates them for repeatable analysis.",
                "e": "RQ1-RQ3 appear in Report Section 1.3; the corresponding results are reported in Section 4.4.",
                "t": "Do not shift the questions into customer satisfaction or sales impact; those outcomes were not evaluated.",
            },
            {
                "q": "Why did you use Design Science Research instead of a conventional machine-learning experiment?",
                "a": "The main research output is a functioning computing artefact, so Design Science fits the objective of building and evaluating a solution in context. I needed to study not only ranking behaviour, but also preference input, ordering, administration, persistence, and explanation. A pure model-comparison experiment would not evaluate whether the recommendation mechanism works inside a usable restaurant workflow. The controlled recommendation experiments remain part of the evaluation, but they are not the whole project.",
                "e": "The report frames the contribution using Design Science and cites Baskerville et al. (2018) on balancing artefact and theory.",
                "t": "Do not imply that building software alone is research; the evaluated design decisions and resulting evidence are what support the research contribution.",
            },
            {
                "q": "Why restrict the work to one restaurant? Does that not destroy generalisability?",
                "a": "It limits generalisability, and I state that directly. The single-restaurant scope was deliberate because the research problem concerns small restaurants with one menu, limited infrastructure, and sparse order history. It allowed me to build and test the full workflow within the available time. The findings demonstrate technical feasibility and behaviour for this dataset; they do not prove that the same parameters will work for other cuisines, menu sizes, or customer populations.",
                "e": "Report Section 1.5 identifies the single-restaurant environment as a research limitation.",
                "t": "Do not say the method is universally suitable for all restaurants.",
            },
            {
                "q": "Why a mobile web application instead of a native app or a tablet at every table?",
                "a": "A mobile web application supports QR access from a customer's own phone, so the restaurant does not need dedicated hardware at every table and the customer does not need to install an app. It also keeps deployment simple because customer and administrator interfaces use one Rust server. The trade-off is that browser sessions and connectivity become dependencies, and the prototype has not been evaluated as an offline-capable or production-scale mobile service.",
                "e": "The report's problem context and architecture describe a mobile-first browser interface for customers and staff.",
                "t": "Do not claim lower cost without a measured cost analysis; describe it as a design rationale.",
            },
            {
                "q": "Where is your evidence that real customers even want this feature?",
                "a": "I do not have a real-customer usability study, so I cannot claim proven demand, higher satisfaction, faster decisions, or increased sales. The motivation is supported conceptually by food-recommender literature and by the problem of uniform digital menus, while this project evaluates technical feasibility and recommendation behaviour. A proper next stage would recruit customers and staff to measure task completion, decision time, recommendation acceptance, perceived relevance, and explanation usefulness.",
                "e": "The report explicitly excludes claims about customer satisfaction, decision time, revenue, and retention.",
                "t": "Never convert a literature-based motivation into a claim of measured customer demand.",
            },
        ],
    ),
    (
        "2. Dataset, Metadata, and Scope",
        [
            {
                "q": "You have only 30 dishes and 60 historical orders. Why should anyone trust the results?",
                "a": "The results should be trusted only within their stated scope. The dataset is sufficient to demonstrate deterministic behaviour, sparse-data handling, order persistence, and the direction of controlled ranking changes. It is not sufficient to estimate general prediction accuracy or population-level preferences. The system therefore reduces reliance on collaborative evidence when it is weak and exposes evidence confidence rather than presenting every score as equally reliable.",
                "e": "The evaluated baseline contains 30 dishes and 60 historical baskets; the report repeatedly labels the dataset as limited.",
                "t": "Do not describe 60 orders as a large or representative dataset.",
            },
            {
                "q": "Are these genuine customer orders, or did you manufacture the data?",
                "a": "The project uses a controlled prototype dataset, not a longitudinal production dataset from a live restaurant. The baseline CSV is treated as historical order evidence for the artefact. RQ2 adds temporary simulated baskets only to an in-memory copy, and the method-comparison cases are drawn from the baseline baskets. I must label those conditions clearly because synthetic and controlled evidence can demonstrate mechanism behaviour but cannot establish real-world customer preference.",
                "e": "Report Sections 1.5 and 3.6 identify controlled and synthetic evaluation as a limitation and describe non-destructive temporary orders.",
                "t": "Do not call the dataset 'real customer behaviour' unless its provenance can be independently verified.",
            },
            {
                "q": "How do you know the ingredients, tags, prices, and categories in your CSV are correct?",
                "a": "The loader validates structure, normalises text, and rejects or reports malformed records, but it cannot prove semantic accuracy. Recommendation quality depends on staff-curated metadata. A production trial would need controlled ingredient names, verified dietary and allergen fields, preparation information, and an ownership process for updating menu records. The current dataset is adequate for a prototype, but metadata quality remains an external dependency.",
                "e": "The data layer uses CSV/Serde models and normalisation; Report Section 5.4 recommends stronger standardised menu metadata.",
                "t": "Do not confuse successful CSV parsing with verified food information.",
            },
            {
                "q": "What happens with chilli versus chili, singular versus plural, or inconsistent ingredient names?",
                "a": "The current cleaning handles trimming, case normalisation, empty values, and exact normalised terms. Search also supports aliases and food concepts, but that does not make the recommendation data a complete food ontology. Unmodelled spelling variants or synonyms can still miss matches. A stronger version would use controlled vocabulary IDs, explicit synonym mapping, and validation during dish entry rather than relying only on free text.",
                "e": "The implemented search and loader normalise text, while the report lists standardised ingredient names as future work.",
                "t": "Do not claim semantic understanding equivalent to a language model or ontology.",
            },
            {
                "q": "How does your system handle a brand-new dish with no order history?",
                "a": "That is an item cold-start case. The new dish can still receive a content score from ingredients and tags, and it can appear through popularity/time fallback rules only after suitable data exist. Its collaborative score starts at zero because there is no co-order evidence. This is one reason the hybrid system does not rely solely on collaborative filtering. The limitation is that weak or incomplete menu metadata will also weaken the cold-start recommendation.",
                "e": "The production pipeline combines content, co-order, popularity, and time signals and adapts weights to the evidence available.",
                "t": "Do not invent collaborative evidence for a new item.",
            },
            {
                "q": "Why are you using CSV and JSONL instead of a proper database?",
                "a": "CSV and JSONL keep the prototype lightweight, transparent, easy to inspect, and easy to demonstrate. They are appropriate for one local process and a small dataset, but they do not provide transactional updates, multi-instance consistency, durable sessions, or robust concurrent access. The architecture isolates file persistence from recommendation and web logic so a relational database can replace it later without rewriting the scoring modules.",
                "e": "The report lists local file-based storage as a limitation and recommends a relational database for a controlled restaurant trial.",
                "t": "Do not call file persistence production-grade simply because restart persistence works.",
            },
        ],
    ),
    (
        "3. Ingredient Logic and Restrictions",
        [
            {
                "q": "State your ingredient score exactly. If you cannot, did you really build it?",
                "a": "For an eligible dish, the ingredient component is the number of matched liked ingredients divided by the dish's total ingredient count, plus 0.15 for each matched preferred tag, capped at 1.0. If a disliked term matches, the dish is removed from recommendation eligibility before ranking. This component is then combined with other signals through the selected recommendation configuration. The formula is intentionally simple and explainable; it is not a calibrated probability of preference.",
                "e": "The implementation in ingredient_filter.rs computes matched_liked / max(total_ingredients, 1) + 0.15 per tag, capped at 1.0.",
                "t": "Do not describe the score as the probability that a customer will like the dish.",
            },
            {
                "q": "Does dividing by total ingredients unfairly punish complex dishes?",
                "a": "It can. A liked ingredient contributes a smaller proportion to a dish with a long ingredient list than to a simpler dish. That is a deliberate baseline interpretation of match density, but it creates a length bias. A future comparison could test raw match count, weighted ingredients, TF-IDF-style rarity, or separate primary and secondary ingredients. I retained the ratio because it is transparent and deterministic, not because it is theoretically optimal.",
                "e": "The denominator is the candidate dish's ingredient count, so this limitation follows directly from the implemented formula.",
                "t": "Do not deny the bias; acknowledge it and explain why the baseline remained useful.",
            },
            {
                "q": "Why is the preferred-tag bonus exactly 0.15? That sounds arbitrary.",
                "a": "It is a heuristic design parameter chosen to make tags a smaller secondary signal than ingredient matching. It is centralised, capped by the component's 1.0 limit, and tested for deterministic behaviour. It was not learned from customer outcomes, so I cannot claim that 0.15 is optimal. A stronger study would perform sensitivity analysis or tune the value using held-out data and user feedback.",
                "e": "The source code documents the 0.15 per-tag bonus as a simple prototype rule.",
                "t": "Do not invent a statistical derivation for the value.",
            },
            {
                "q": "What if the same ingredient is selected as both liked and disliked?",
                "a": "The preference interface prevents the conflict by removing an ingredient from the opposite set when it is selected. At the recommendation boundary, disliked matching has priority because eligibility is checked before scoring. Therefore, even if inconsistent input reached the engine, the safer deterministic interpretation is exclusion rather than allowing the liked score to override the restriction.",
                "e": "The UI state enforces mutual exclusivity, and the ingredient scorer returns zero immediately when a disliked match is present.",
                "t": "Do not leave the conflict resolution ambiguous; state that the restriction wins.",
            },
            {
                "q": "You use hard exclusions. Are you claiming this system is safe for food allergies?",
                "a": "No. The disliked-ingredient feature is preference support, not certified allergy management. It depends on menu metadata and does not model hidden ingredients, substitutions, kitchen practices, or cross-contamination. A production allergy feature would require a controlled allergen taxonomy, restaurant verification, governance procedures, prominent warnings, and legal or regulatory review. I would never advise a customer to rely on this prototype for medical safety.",
                "e": "Certified allergy and nutritional assessment are explicitly outside the project scope.",
                "t": "Never use the phrases 'allergy-safe' or 'guaranteed safe.'",
            },
            {
                "q": "Can popularity or co-order evidence force a disliked dish back into the list?",
                "a": "No. Eligibility is resolved before ranking. A dish that is unavailable, already selected where exclusion is required, or matched to a disliked term is removed from the candidate set. Collaborative, popularity, time, and diversity signals operate only on eligible candidates. This separation is important because a restriction must not be traded against a high behavioural score.",
                "e": "The report identifies eligibility, ranking, and evidence confidence as separate recommendation decisions.",
                "t": "Do not say a large hybrid score can compensate for a hard restriction.",
            },
        ],
    ),
    (
        "4. Collaborative, Hybrid, and Explainability Logic",
        [
            {
                "q": "You call this collaborative filtering, but where are the user-user similarities and ratings?",
                "a": "This is item-to-item collaborative evidence derived from baskets, not user-user rating prediction. The matrix counts which dishes appear together in completed orders. The current selected dishes provide the request context, and related candidates receive a co-order score. This approach fits the dataset because the restaurant has basket history but little repeated identity or rating data. Personalisation comes from combining current-session preferences and context, not from a permanent user profile.",
                "e": "Research Objective 3 and the implementation both describe an item-to-item co-order component without explicit ratings.",
                "t": "Do not describe it as user-based collaborative filtering.",
            },
            {
                "q": "If two dishes are ordered together, why assume one causes preference for the other?",
                "a": "I do not assume causation. Co-ordering is an association signal: it says the dishes have appeared in the same baskets more often, not that one causes purchase of the other or that every customer will like both. The explanation therefore uses wording such as 'often ordered with,' not 'you will like this.' Support, confidence, lift, and evidence confidence help describe the strength and context of the association.",
                "e": "The report evaluates association measures and explicitly treats the recommender as decision support.",
                "t": "Avoid causal language such as 'D01 makes customers buy D07.'",
            },
            {
                "q": "What stops one duplicated dish ID inside an order from inflating the pair count?",
                "a": "The co-order builder first places the order's dish IDs in a set. It then counts each unique pair once for that basket and stores the relationship symmetrically in both directions. Therefore, a repeated ID in one row does not create multiple pair observations. Quantities are intentionally ignored in this prototype; the evidence unit is whether two dish IDs co-occurred in a basket.",
                "e": "collaborative_filter.rs uses a HashSet before generating pairs and documents the duplicate-prevention rule.",
                "t": "Do not claim quantity-sensitive collaborative filtering; it is basket presence/absence.",
            },
            {
                "q": "How is the co-order score normalised, and what does 1.0 actually mean?",
                "a": "For the current selected-dish context, the candidate's related counts are compared with the strongest related candidate count in that same request. The resulting score is bounded between 0 and 1. A score of 1.0 means the candidate is strongest relative to the other candidates for that context; it does not mean a 100 percent purchase probability. Request-scoped normalisation keeps candidates comparable within one ranking but not necessarily across unrelated requests.",
                "e": "The collaborative module computes strongest_related_count once per request and reuses it as the denominator.",
                "t": "Do not interpret the normalised score as calibrated probability or confidence.",
            },
            {
                "q": "Explain support, confidence, and lift without hiding behind jargon.",
                "a": "Support is the share of all baskets containing the pair. Directional confidence asks: among baskets containing the anchor dish, what share also contains the candidate? Lift compares that confidence with how common the candidate is overall. Lift above 1 indicates the pair occurs together more than the candidate's baseline prevalence would suggest. These are descriptive association measures, and small counts can still produce unstable-looking values.",
                "e": "RQ2 reports pair count, support, directional confidence, lift, and rank together rather than relying on one metric.",
                "t": "Do not present high lift from a tiny sample as strong causal evidence.",
            },
            {
                "q": "A rare pair can produce impressive lift. How do you stop the system overreacting?",
                "a": "The system separates ranking score from evidence confidence and adapts the collaborative contribution using dataset size, selected-context frequency, and pair strength. This reduces the influence of sparse behavioural evidence, but it does not remove statistical uncertainty. The confidence formula is a heuristic evidence indicator, so a production version should add minimum-support rules, recency weighting, uncertainty estimates, and validation on more data.",
                "e": "The adaptive confidence combines dataset, context, and pair-strength terms; the report warns that evidence confidence is not predictive probability.",
                "t": "Do not say the heuristic makes rare-pair estimates statistically reliable.",
            },
            {
                "q": "Why include popularity and time context? Are they not just noise?",
                "a": "They provide deterministic fallback signals when explicit preferences or co-order context are weak. Popularity can prevent an empty or arbitrary list, while time context can apply a small restaurant-relevant prior. They are not treated as proof of individual taste. Their influence is controlled by the adaptive weighting, and disliked restrictions still apply first. A future study should test whether the time rules improve user outcomes rather than assuming they do.",
                "e": "The production model includes content, co-order, popularity, and time components and changes their weights based on available evidence.",
                "t": "Do not call a time-of-day rule learned personalisation.",
            },
            {
                "q": "Your adaptive weights and thresholds look like magic numbers. Where did they come from?",
                "a": "They are engineering heuristics for a transparent prototype, not parameters learned from outcomes. The implementation centralises the thresholds, keeps the weight changes monotonic, normalises each final weight set to 1.0, and tests the regimes. For example, confidence considers targets for total orders, context orders, and pair counts. I can defend their internal consistency, but not claim that they are optimal for another restaurant. Sensitivity analysis and held-out tuning are future work.",
                "e": "adaptive.rs defines explicit targets and tests every evidence regime; the report labels the system deterministic and lightweight.",
                "t": "Do not pretend the thresholds were statistically estimated.",
            },
            {
                "q": "Your report compares a fixed 0.4/0.6 hybrid, but the application uses adaptive weights. Is that a contradiction?",
                "a": "No, because they serve different purposes. The fixed 0.4 content and 0.6 co-order configuration is an experimental comparator used to make RQ3 repeatable and isolate method behaviour. The production recommender adapts weights to the available preference and behavioural evidence. I should label this distinction clearly: the experiment compares controlled methods; it does not claim that the fixed configuration is the deployed policy.",
                "e": "The Recommendation Tester exposes ingredient-only, co-order-only, and fixed 0.4/0.6 modes, while the customer path uses adaptive scoring.",
                "t": "Do not quote 0.4/0.6 as the universal production formula.",
            },
            {
                "q": "Your evidence-confidence badge looks scientific. Is it a probability that the recommendation is correct?",
                "a": "No. It is a bounded indicator of how much supporting evidence is available, based on factors such as dataset size, context frequency, pair strength, preference coverage, and popularity. It is deliberately separate from the ranking score. A high evidence label does not mean the customer has a corresponding probability of liking or buying the dish, and it has not been statistically calibrated.",
                "e": "The report states explicitly that evidence confidence is not predictive probability.",
                "t": "Never translate a confidence label into a percentage chance of customer acceptance.",
            },
            {
                "q": "Are your explanations genuinely faithful, or are they decorative text added after ranking?",
                "a": "They are generated from the same matched ingredients, matched tags, selected dish IDs, component scores, and exclusion checks used by the ranking pipeline. That makes them faithful to the implemented rules. However, faithfulness is not the same as usefulness: the project did not run a user study to test whether customers understand, trust, or act on the explanations. That remains a separate evaluation requirement.",
                "e": "The ingredient and hybrid modules build explanations from computed match sets and co-order influences; Zhang and Chen (2020) support evaluating explainability explicitly.",
                "t": "Do not claim that explanations improved trust; that outcome was not measured.",
            },
        ],
    ),
    (
        "5. Evaluation Design and Results",
        [
            {
                "q": "Five hidden-dish cases are nowhere near enough. Why did you report them?",
                "a": "They are a small controlled demonstration, not a general benchmark. I reported them because they make the method behaviour observable and repeatable within the project dataset. The correct interpretation is 'in these five selected cases,' not 'the hybrid is 100 percent accurate.' A stronger evaluation would use many more baskets, multiple preference profiles, a held-out protocol, and uncertainty estimates.",
                "e": "The report limits RQ3 conclusions to the 30-dish menu, 60-order baseline, and five selected scenarios.",
                "t": "Do not generalise five cases to all customers or future orders.",
            },
            {
                "q": "You left the test basket in the historical data. Is that not data leakage?",
                "a": "Yes, it can make the recovery result optimistic. The experiment hides one dish from the input context but retains the baseline basket in the historical evidence, so it demonstrates retrieval under known historical relationships rather than strict out-of-sample prediction. I disclose this limitation. The stronger next protocol is leave-one-basket-out or chronological evaluation, where the entire test basket is removed before the hidden item is predicted.",
                "e": "Report Section 4.4.5 and Future Work explicitly identify the retained baskets and recommend leave-one-basket-out or chronological splits; Ji et al. (2023) is cited on leakage.",
                "t": "Do not deny leakage or describe this as an independent train-test result.",
            },
            {
                "q": "You reported 100 percent Hit@3. Is that not a misleading headline?",
                "a": "It would be misleading without the denominator and conditions. The exact result is 5 successful cases out of 5 selected historical baskets for co-order-only and fixed hybrid, with the baskets retained in the historical baseline. Hit@3 is also not classification accuracy. I present it as controlled evidence that co-order relationships were useful in those cases, not proof of 100 percent predictive performance.",
                "e": "RQ3 reports 5/5, 100 percent Hit@3, and average hidden-dish rank 1.80, alongside explicit scope limitations.",
                "t": "Always say '5 out of 5 selected cases,' not merely '100 percent accuracy.'",
            },
            {
                "q": "Why did ingredient-only perform so badly at 20 percent Hit@3? Does that mean your content model failed?",
                "a": "It means the fixed liked preferences, rice and chicken, often did not describe the hidden dishes in those five baskets. Ingredient-only had no behavioural signal to recover a complementary item whose ingredients differed from those preferences. This exposes a real limitation of content-only ranking and helps explain why co-order evidence can be useful. It does not prove that ingredient filtering is generally poor, because the cases and preference profile were very limited.",
                "e": "Ingredient-only recovered 1 of 5 targets with average rank 14.20 under one fixed preference profile.",
                "t": "Do not call the 20 percent result a universal content-based accuracy rate.",
            },
            {
                "q": "Using the same 'rice and chicken' preferences in every case seems biased. Why did you do that?",
                "a": "Keeping the preference profile fixed controls one source of variation, so differences among cases are easier to interpret. However, the choice can advantage or disadvantage ingredient-only depending on the hidden dish and does not represent varied users. A stronger study should repeat the experiment with multiple predefined profiles, no-preference conditions, and profiles aligned and misaligned with each hidden target.",
                "e": "The method-comparison protocol states that liked ingredients rice and chicken were held constant across all five cases.",
                "t": "Do not claim that one preference profile represents customer diversity.",
            },
            {
                "q": "Why use Hit@3 and hidden-dish rank instead of precision, recall, NDCG, or MRR?",
                "a": "Each case has one hidden target, and the customer interface emphasises a short recommendation list. Hit@3 answers whether that target is visible near the top, while hidden-dish rank preserves its exact position. Those metrics are easy to explain for a prototype. They are not a complete evaluation. With more cases, I would add Mean Reciprocal Rank, NDCG, catalogue coverage, diversity, novelty, restriction violations, and confidence intervals.",
                "e": "Report Section 3.6 defines Hit@K and hidden-dish rank and explains why they are interpreted together.",
                "t": "Do not imply that Hit@3 measures all aspects of recommendation quality.",
            },
            {
                "q": "Where are your statistical significance tests and confidence intervals?",
                "a": "They are not present because five cases do not support meaningful inferential claims. The analysis is descriptive and mechanism-focused. Reporting a p-value here would give a false impression of statistical strength. A future evaluation needs a larger, pre-specified sample, appropriate paired tests or bootstrap intervals, and a protocol that removes leakage before inferential comparison.",
                "e": "The report describes the results as controlled and limited rather than statistically generalisable.",
                "t": "Do not invent significance or use 'significant' as a synonym for 'large.'",
            },
            {
                "q": "Why did you not compare against random and popularity-only baselines?",
                "a": "That is a gap in the RQ3 comparison. Ingredient-only, co-order-only, and fixed hybrid answer the stated method question, but random and popularity-only baselines would show whether the observed ranks exceed simple alternatives. The production system does include popularity as a fallback, yet it was not isolated in that result table. Adding those baselines is a clear improvement to the evaluation design.",
                "e": "The reported method table contains only the three specified modes; popularity exists in production but is not a separate RQ3 comparator.",
                "t": "Do not suggest that outperforming a popularity baseline was demonstrated.",
            },
            {
                "q": "Your ingredient-impact evidence looks cherry-picked. How many cases did you actually test?",
                "a": "The reported RQ1 table contains three controlled cases: liking bean sprouts moved Char Kway Teow from rank 10 to 1; disliking banana excluded Pisang Goreng from rank 1; and liking rice while disliking banana moved Ketupat from rank 2 to 1 while preserving the exclusion. These cases demonstrate direction and rule compliance, but they are illustrative, not a representative sample of all ingredients and dishes.",
                "e": "The RQ1 result file and Report Section 4.4.1 contain three predefined cases.",
                "t": "Do not imply that every possible preference combination was empirically evaluated.",
            },
            {
                "q": "By adding temporary co-orders, are you merely forcing the answer you wanted?",
                "a": "The intervention is intentionally controlled: the purpose of RQ2 is to test whether the system responds when co-order evidence increases. All other conditions remain fixed, and the temporary baskets are applied to an in-memory copy so operational history is unchanged. For D01 and D07, adding three temporary baskets increased the pair evidence and moved D07 from rank 7 to 1. This demonstrates sensitivity to the intended signal, not realism of future customer behaviour.",
                "e": "RQ2 records baseline, +3, and +5 basket conditions and reports pair count, support, confidence, lift, and rank.",
                "t": "Do not present simulated baskets as observed market behaviour.",
            },
        ],
    ),
    (
        "6. Testing, Architecture, and User Experience",
        [
            {
                "q": "You claim 115 Rust tests and 55 report checks. Are you double-counting to make the project look stronger?",
                "a": "They are different scopes and should not be added as 170 independent tests. The 115 Rust tests exercise implementation units and integration behaviour in the codebase. The 55 selected report-level checks summarise unit, integration, end-to-end, security/access, and responsive inspection evidence for the report. Both sets passed, but overlap exists. I report them separately and do not treat the total as a statistical quality score.",
                "e": "System Testing Results states 115 automated Rust tests and 55 selected report-level checks, with their purposes described separately.",
                "t": "Do not say '170 unique tests prove the system has no bugs.'",
            },
            {
                "q": "All tests passed, so are you claiming the system is reliable and secure?",
                "a": "No. Passing tests means the specified cases behaved as expected in the tested environment. It does not prove absence of defects, security vulnerabilities, performance bottlenecks, or usability problems. The security checks verify session separation and protected routes, but they are not penetration testing or certification. Reliability under crashes, multiple server instances, and production load remains outside the evidence.",
                "e": "The report explicitly says the security checks are prototype verification, not security certification.",
                "t": "Never equate test pass rate with proof of production readiness.",
            },
            {
                "q": "Where is your usability study? Screenshots are not usability evidence.",
                "a": "That criticism is correct. Responsive inspection and end-to-end workflow checks show that controls render and function at 390 x 844, 768 x 1024, and 1440 x 900, but they do not measure comprehension, effort, satisfaction, or accessibility with real participants. A follow-up study should use task scenarios, completion rates, time-on-task, errors, a usability questionnaire, and interviews with customers and restaurant staff.",
                "e": "The report separates responsive-interface inspection from the future need for real-customer and staff evaluation.",
                "t": "Do not call visual QA a user study.",
            },
            {
                "q": "Why Rust and Axum? Did you measure that they were better than Node.js or Python?",
                "a": "I selected Rust and Axum for strong typing, memory safety, predictable performance, and a lightweight single-server architecture. They also let the recommendation and web layers share typed models without a separate service. I did not benchmark alternative stacks, so I cannot claim Rust was objectively faster or cheaper for this project. It was a technically suitable implementation choice, not an experimental finding.",
                "e": "The implemented stack is Rust, Axum, Tokio, server-rendered HTML, JavaScript, and CSS.",
                "t": "Do not turn a technology preference into an unsupported performance comparison.",
            },
            {
                "q": "Why server-rendered HTML and plain JavaScript instead of React or another modern frontend?",
                "a": "The project prioritises a lightweight prototype with few dependencies and clear module boundaries. Server rendering simplifies initial page delivery and keeps routing and state decisions on the Rust side, while focused JavaScript handles search suggestions, scrolling, cart interactions, and mobile gestures. The trade-off is less client-side structure for very complex interactions. For this scope, avoiding a large frontend framework reduced deployment and maintenance complexity.",
                "e": "The architecture deliberately separates Rust handlers/views from static CSS and JavaScript rather than placing everything in main.rs.",
                "t": "Do not claim that a framework would always be worse; explain the scope-dependent trade-off.",
            },
            {
                "q": "Why does search not filter the Menu? Is that not poor usability?",
                "a": "The design separates three responsibilities: Smart Search locates a dish, Recommended for You ranks suggestions, and Menu preserves the full available catalogue. Search suggestions can scroll to and highlight a matching card without hiding or reordering other dishes. This maintains customer choice and prevents search or recommendation state from making dishes appear unavailable. It is a deliberate information-architecture decision, although a real usability study should compare it with conventional filtering.",
                "e": "System test UT-05 confirms that searching 'laksa' returns three suggestions while the static Menu remains at 30 dishes.",
                "t": "Do not claim the design is proven superior; it is an intentional, tested behaviour.",
            },
            {
                "q": "Diversity reranking can lower relevance. Why interfere with the score order?",
                "a": "The system exposes Familiar, Balanced, and Discover modes so diversity is a user-facing trade-off rather than a hidden replacement of relevance. Reranking can reduce near-duplicate categories and broaden menu discovery, but it may move a slightly lower-scoring dish upward. Restrictions remain non-negotiable. The current project verifies deterministic behaviour; it does not establish which diversity mode users prefer or whether diversity improves satisfaction.",
                "e": "Diversity modes are listed in the project scope and operate after eligibility and base scoring.",
                "t": "Do not describe diversity as an automatic improvement for every customer.",
            },
            {
                "q": "Does your budget-aware meal set guarantee a nutritionally valid or optimally priced meal?",
                "a": "No. It is a prototype set generator that respects the available menu prices, selected budget, availability, and preference restrictions while producing a practical combination. It is not a nutrition optimiser, payment quote, or proof of global mathematical optimality. Its value is demonstrating how recommendation constraints can support a complete ordering decision. Production use would require verified prices, tax/service-charge rules, portions, and nutritional data.",
                "e": "Budget-aware meal-set generation is in scope, while online payment and certified nutritional assessment are explicitly excluded.",
                "t": "Do not call the generated set nutritionally balanced or financially final.",
            },
        ],
    ),
    (
        "7. Security, Ethics, Persistence, and Deployment",
        [
            {
                "q": "Your order history is stored in files. What happens under concurrent checkout or a crash halfway through a write?",
                "a": "The prototype is designed for one local server process and validates that a completed order is appended once, loaded into memory immediately, and survives restart. That does not provide database-grade atomicity or multi-instance concurrency. A crash during a file operation or simultaneous writers remains a risk. A production pilot should use a relational database with transactions, constraints, backups, and recovery procedures.",
                "e": "End-to-end testing confirmed append-once persistence and reload, while the report recommends replacing local files for operational deployment.",
                "t": "Do not confuse duplicate prevention in one process with full transactional integrity.",
            },
            {
                "q": "Your administrator username and password are both 'admin.' Why should anyone take security seriously?",
                "a": "Those are deliberately documented local-demo defaults, not acceptable production credentials. The server supports environment-variable overrides, uses separate customer and administrator sessions, HttpOnly cookies, SameSite policy, and protects administrator mutation routes. However, the prototype does not implement persistent staff accounts, password hashing, role-based access, CSRF tokens, rate limiting, or TLS termination. I would require those controls before any real deployment.",
                "e": "README and architecture documentation identify admin/admin as a demo default and list the missing production controls.",
                "t": "Do not defend admin/admin as secure; defend only its controlled demo purpose and the documented boundary.",
            },
            {
                "q": "What customer data do you collect, and how is privacy protected?",
                "a": "The prototype uses temporary restaurant-session information and does not build permanent cross-visit customer profiles. Customer and administrator sessions are separated, and customers are scoped to their own order view. This data-minimisation approach reduces long-term profiling, but order logs and session data still require access control, retention rules, and secure transport in production. The project does not claim formal privacy compliance assessment.",
                "e": "Permanent cross-visit profiles are out of scope; the architecture uses separate in-memory session maps and cookies.",
                "t": "Do not say 'we collect no personal data' unless every registration and log field supports that claim.",
            },
            {
                "q": "Could fake or manipulated orders poison your recommender?",
                "a": "Yes. Because completed baskets update popularity and co-order evidence, coordinated or erroneous orders could influence future rankings. The prototype reduces accidental duplication by appending a completed order once, but it has no fraud detection, trust weighting, or anomaly monitoring. Production safeguards could require verified payment or staff confirmation, minimum support, recency decay, outlier detection, audit logs, and limits on how quickly evidence can change weights.",
                "e": "The closed feedback loop intentionally uses completed orders as future evidence; manipulation resistance is not part of the current evaluation.",
                "t": "Do not claim the feedback loop is automatically beneficial or tamper-proof.",
            },
            {
                "q": "Does your recommender create a popularity feedback loop that keeps recommending the same dishes?",
                "a": "That risk exists. Popular items can receive more exposure, generate more orders, and become even more popular. The system reduces this through explicit preferences, co-order context, evidence-aware weighting, and optional diversity modes, but those controls do not eliminate feedback bias. A production study should monitor catalogue coverage, exposure distribution, novelty, and long-term concentration, and could add time decay or controlled exploration.",
                "e": "Popularity is a fallback component and diversity modes are available, but long-term fairness and exposure were not evaluated.",
                "t": "Do not imply that diversity reranking fully solves popularity bias.",
            },
            {
                "q": "Would you deploy this in a real restaurant tomorrow? If not, what have you actually achieved?",
                "a": "I would not deploy it as a production system tomorrow. I have achieved a functioning and evaluated prototype that demonstrates technical feasibility: mobile ordering, transparent recommendation, cart and status workflow, administration, persistence, and controlled experiments operate together. A controlled pilot would still require a database, hardened authentication and HTTPS, verified menu/allergen data, concurrency and load testing, operational backup, staff training, and real-user evaluation. The next research step is a leakage-free offline study on more genuine orders, followed by a small supervised user trial.",
                "e": "The report conclusion explicitly frames the result as technical feasibility rather than production readiness or commercial effectiveness.",
                "t": "A strong defence is a precise boundary, not pretending the prototype is deployment-ready.",
            },
        ],
    ),
]


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = "w:" + edge
        node = tc_mar.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_paragraph_border(paragraph, side, color, size=8, space=4):
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border = p_bdr.find(qn(f"w:{side}"))
    if border is None:
        border = OxmlElement(f"w:{side}")
        p_bdr.append(border)
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color)


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_custom_numbering(document, kind):
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
        if node.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId")) is not None
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    p_pr.append(indent)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_pr.append(ilvl)
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(num_id_node)
    p_pr.append(num_pr)


def add_list_item(document, text, num_id):
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    apply_num(p, num_id)
    run = p.add_run(text)
    set_run_font(run, size=11, color=NAVY)
    return p


def configure_styles(document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.append(fld)


def configure_section(section, first=False):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = first

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), alignment=2)
    left = p.add_run("PERSONALIZED RESTAURANT ORDERING SYSTEM")
    set_run_font(left, size=8.5, color=MUTED, bold=True)
    p.add_run("\t")
    right = p.add_run("MOCK VIVA GUIDE")
    set_run_font(right, size=8.5, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    label = fp.add_run("FYP DEFENCE PREPARATION  |  ")
    set_run_font(label, size=8.5, color=MUTED)
    add_page_number(fp)


def add_inline_label_paragraph(document, label, text, label_color=ORANGE):
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    lr = p.add_run(label + " ")
    set_run_font(lr, size=11, color=label_color, bold=True)
    tr = p.add_run(text)
    set_run_font(tr, size=11, color=NAVY)
    return p


def add_evidence(document, text):
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.14)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    set_paragraph_shading(p, LIGHT_GRAY)
    set_paragraph_border(p, "left", BLUE, size=12, space=5)
    label = p.add_run("Evidence anchor: ")
    set_run_font(label, size=9.5, color=DARK_BLUE, bold=True)
    value = p.add_run(text)
    set_run_font(value, size=9.5, color=MUTED)
    return p


def add_trap(document, text):
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.14)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    set_paragraph_shading(p, LIGHT_RED)
    set_paragraph_border(p, "left", RISK_RED, size=12, space=5)
    label = p.add_run("Do not overclaim: ")
    set_run_font(label, size=9.5, color=RISK_RED, bold=True)
    value = p.add_run(text)
    set_run_font(value, size=9.5, color=RISK_RED, italic=True)
    return p


def add_cover(document):
    for _ in range(4):
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(8)

    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    r = kicker.add_run("FYP VIVA DEFENCE GUIDE")
    set_run_font(r, size=11, color=ORANGE, bold=True)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    r = title.add_run("Adversarial Mock Viva")
    set_run_font(r, size=30, color=NAVY, bold=True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(4)
    r = subtitle.add_run("54 Tough Questions and Defensible Answers")
    set_run_font(r, size=15, color=DARK_BLUE, bold=True)

    project = document.add_paragraph()
    project.alignment = WD_ALIGN_PARAGRAPH.CENTER
    project.paragraph_format.space_after = Pt(28)
    r = project.add_run("Personalized Restaurant Ordering System")
    set_run_font(r, size=14, color=BLUE)

    statement = document.add_paragraph()
    statement.alignment = WD_ALIGN_PARAGRAPH.CENTER
    statement.paragraph_format.left_indent = Inches(0.7)
    statement.paragraph_format.right_indent = Inches(0.7)
    statement.paragraph_format.space_after = Pt(58)
    statement.paragraph_format.line_spacing = 1.25
    r = statement.add_run(
        "The examiner voice is intentionally severe. The model answers remain calm, precise, and honest about the prototype's limits."
    )
    set_run_font(r, size=11, color=MUTED, italic=True)

    author = document.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.paragraph_format.space_after = Pt(4)
    r = author.add_run("Prepared for Yeap Chan Leong (22049837)")
    set_run_font(r, size=11, color=NAVY, bold=True)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(4)
    r = meta.add_run("Bachelor of Science (Honours) in Computer Science | August 2026")
    set_run_font(r, size=9.5, color=MUTED)

    document.add_page_break()


def add_front_matter(document, bullet_num, decimal_num):
    document.add_heading("How to Use This Guide", level=1)
    intro = document.add_paragraph(
        "Do not memorise every sentence. Practise the first sentence of each answer, then use the evidence anchor and limitation only if the examiner presses further. A strong viva answer is usually 20-40 seconds."
    )
    intro.paragraph_format.space_after = Pt(8)

    for item in (
        "Answer the exact question first. Do not begin with background that avoids the challenge.",
        "Give one concrete implementation or result as evidence.",
        "State the boundary honestly when the evidence is weak or the feature is prototype-level.",
        "Finish with the specific next step that would strengthen the work.",
    ):
        add_list_item(document, item, decimal_num)

    document.add_heading("Evidence Snapshot to Memorise", level=1)
    data = [
        ("Scope", "Single-restaurant, mobile-first Rust/Axum web prototype."),
        ("Baseline data", "30 dishes and 60 initial historical order baskets."),
        ("System checks", "115 Rust tests passed; 55 selected report-level checks passed. Do not add these as 170 unique tests."),
        ("RQ1", "Bean sprouts: D14 rank 10 to 1. Banana disliked: D30 rank 1 to Excluded."),
        ("RQ2", "For D01 with D07, three temporary co-orders moved D07 from rank 7 to 1."),
        ("RQ3", "Five cases: ingredient-only Hit@3 20%, average rank 14.20; co-order-only and fixed 0.4/0.6 hybrid Hit@3 100%, average rank 1.80."),
        ("Critical limitation", "The five test baskets remained in historical evidence, so RQ3 is controlled retrieval, not leakage-free prediction."),
        ("Claim boundary", "Technical feasibility and deterministic behaviour, not production readiness, commercial impact, or allergy safety."),
    ]
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in data:
        cells = table.add_row().cells
        set_cell_shading(cells[0], LIGHT_BLUE)
        set_cell_shading(cells[1], WHITE)
        p0 = cells[0].paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        r0 = p0.add_run(label)
        set_run_font(r0, size=10, color=DARK_BLUE, bold=True)
        p1 = cells[1].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        p1.paragraph_format.line_spacing = 1.15
        r1 = p1.add_run(value)
        set_run_font(r1, size=10, color=NAVY)
    set_table_geometry(table, [1701, 7659])

    document.add_heading("Language That Will Protect Your Marks", level=1)
    for item in (
        "Say: 'In these five selected cases...' rather than 'The model is 100 percent accurate.'",
        "Say: 'The result demonstrates technical feasibility...' rather than 'The system is production-ready.'",
        "Say: 'Item-to-item co-order association...' rather than 'The system knows what every user likes.'",
        "Say: 'Evidence confidence is a heuristic indicator...' rather than 'There is an 80 percent chance this recommendation is correct.'",
        "Say: 'Disliked-ingredient preference...' rather than 'Certified allergy protection.'",
        "Say: 'Application-focused contribution...' rather than 'A new AI algorithm.'",
    ):
        add_list_item(document, item, bullet_num)

def add_qa(document):
    question_number = 1
    for section_title, questions in SECTIONS:
        document.add_heading(section_title, level=1)
        for item in questions:
            heading = document.add_heading(f"Q{question_number}. {item['q']}", level=2)
            heading.paragraph_format.keep_with_next = True
            answer = add_inline_label_paragraph(document, "Defensible answer:", item["a"])
            answer.paragraph_format.keep_with_next = False
            add_evidence(document, item["e"])
            add_trap(document, item["t"])
            question_number += 1
    assert question_number == 55, f"Expected 54 questions, found {question_number - 1}"


def add_final_pages(document, bullet_num):
    document.add_heading("Final Defence Checklist", level=1)
    for item in (
        "I can explain the exact ingredient formula without looking at the code.",
        "I can distinguish item-to-item co-ordering from user-based collaborative filtering.",
        "I can explain why fixed 0.4/0.6 is an experiment while production scoring is adaptive.",
        "I will state that 100 percent means 5/5 selected cases, not general accuracy.",
        "I will acknowledge retained-basket leakage and propose leave-one-basket-out or chronological testing.",
        "I will not describe disliked ingredients as certified allergy protection.",
        "I will not describe admin/admin, CSV persistence, or the current security controls as production-ready.",
        "I will distinguish functional correctness from usability, commercial impact, and security certification.",
        "I can demonstrate one complete customer-to-admin order lifecycle if the live demo works.",
        "If the demo fails, I can continue with verified screenshots, test evidence, and a short explanation of the intended flow.",
    ):
        add_list_item(document, item, bullet_num)

    document.add_heading("Fast Rescue Answer", level=1)
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.25
    set_paragraph_shading(p, LIGHT_BLUE)
    set_paragraph_border(p, "left", ORANGE, size=18, space=6)
    r = p.add_run(
        "The precise claim is that I designed and evaluated a transparent, lightweight prototype for one limited-data restaurant. The controlled results show that the implemented preference and co-order signals change rankings as intended. They do not establish production readiness, general prediction accuracy, customer satisfaction, or commercial impact."
    )
    set_run_font(r, size=11.5, color=NAVY, bold=True)

    document.add_heading("Evidence Sources Used for This Guide", level=1)
    for item in (
        "FR_22049837.pdf, especially Sections 1.3-1.5, 3.3.6-3.6, 4.3-4.5, and 5.2-5.5.",
        "output/ppt/personalized-restaurant-ordering-system/Full-Presentation-Script-EN.md.",
        "docs/report-evidence/2026-07-29/SYSTEM_TESTING_RESULTS.md.",
        "docs/report-evidence/2026-07-29/section-4-4/RECOMMENDATION_EVALUATION_RESULTS.md and its CSV result files.",
        "src/recommender/ingredient_filter.rs, collaborative_filter.rs, adaptive.rs, and hybrid.rs.",
        "README.md and docs/architecture.md for deployment and security boundaries.",
    ):
        add_list_item(document, item, bullet_num)

    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(8)
    note.paragraph_format.space_after = Pt(0)
    run = note.add_run(
        "Preparation note: Check the final submitted report and live repository before the viva. If a number or feature changed after this guide was generated, use the submitted evidence as the authority."
    )
    set_run_font(run, size=9.5, color=MUTED, italic=True)


def build_document():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.core_properties.title = "Adversarial Mock Viva - Personalized Restaurant Ordering System"
    document.core_properties.subject = "54 tough FYP viva questions with evidence-grounded answers"
    document.core_properties.author = "Yeap Chan Leong"
    document.core_properties.keywords = "FYP, viva, restaurant recommender, Rust, hybrid recommendation"

    configure_styles(document)
    configure_section(document.sections[0], first=True)
    bullet_num = add_custom_numbering(document, "bullet")
    decimal_num = add_custom_numbering(document, "decimal")

    add_cover(document)
    add_front_matter(document, bullet_num, decimal_num)
    add_qa(document)
    add_final_pages(document, bullet_num)

    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
